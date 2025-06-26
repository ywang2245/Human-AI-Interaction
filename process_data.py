import os
import json
from docx import Document
import re
from typing import Optional

def extract_task_type(filename):
    # Extract task type from filename
    match = re.search(r'Task\s*([ABC])', filename, re.IGNORECASE)
    if match:
        return match.group(1).upper()
    return None

def extract_student_id(filename):
    # Extract student ID from filename
    match = re.search(r'ID(\d+)', filename)
    if match:
        return match.group(1)
    return None

def parse_conversation_file(doc_path):
    """
    The definitive parser. This handles multiple classes of formatting errors
    by pre-processing the text and then using a robust state machine.
    """
    try:
        doc = Document(doc_path)
        all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        # This can happen if the docx file is corrupted
        print(f"ERROR: Could not read {os.path.basename(doc_path)}. Error: {e}")
        return []

    # 1. Pre-filter junk lines copied from web interfaces
    paras = [p for p in all_paras if p.lower() not in ('top of form', 'bottom of form', 'sources')]

    # 2. Pre-process to fix files where turns are pasted into a single paragraph (like ID 67)
    temp_paras = []
    for p in paras:
        # This regex finds "ChatGPT said:" (case-insensitive) that follows "You said:" in the same line
        if p.lower().startswith("you said:") and "chatgpt said:" in p.lower():
            # Split the paragraph into two distinct turns
            parts = re.split(r'(?i)(ChatGPT said:)', p, maxsplit=1)
            temp_paras.append(parts[0].strip())  # "You said: ..."
            temp_paras.append(parts[1] + parts[2]) # "ChatGPT said: ..."
        else:
            temp_paras.append(p)
    paras = temp_paras

    # 3. The robust state machine
    dialogue_history = []
    student_text = []
    gpt_text = []
    # 0: seeking first turn, 1: in student turn, 2: in gpt turn
    current_state = 0
    
    for p in paras:
        is_student_marker = (p.lower() == "you said:")
        is_gpt_marker = (p.lower() == "chatgpt said:")

        if is_student_marker:
            if current_state == 2:  # We were in a GPT turn, so this completes a round
                if student_text or gpt_text:
                    dialogue_history.append({
                        "student_prompt": "\\n".join(student_text),
                        "gpt_response": "\\n".join(gpt_text)
                    })
                student_text, gpt_text = [], []
            
            elif current_state == 1:  # Two student markers in a row
                 if student_text:
                      dialogue_history.append({
                        "student_prompt": "\\n".join(student_text),
                        "gpt_response": ""
                    })
                 student_text = []

            current_state = 1

        elif is_gpt_marker:
            if current_state == 1 and not student_text: # Handle empty student turn
                student_text.append("") # Add an empty placeholder
            current_state = 2

        else:  # It's content, not a marker
            if current_state == 1:
                student_text.append(p)
            elif current_state == 2:
                gpt_text.append(p)

    # After the loop, save the very last turn
    if student_text or gpt_text:
        dialogue_history.append({
            "student_prompt": "\\n".join(student_text),
            "gpt_response": "\\n".join(gpt_text)
        })

    # Add round numbers
    final_history = []
    for i, turn in enumerate(dialogue_history):
        # Ensure we don't add empty rounds
        if turn["student_prompt"] or turn["gpt_response"]:
            turn["round"] = len(final_history) + 1
            final_history.append(turn)

    return final_history

def process_answer_file(file_path):
    doc = Document(file_path)
    # Assuming the final answer is in the last paragraph
    final_submission = doc.paragraphs[-1].text.strip() if doc.paragraphs else ""
    return final_submission

def main():
    # Use relative paths to ensure the script is portable
    script_dir = os.path.dirname(os.path.abspath(__file__))
    base_dir = os.path.dirname(script_dir)  # Go up one level from 'src'

    answers_dir = os.path.join(base_dir, "Answers - Cleaned")
    conversations_dir = os.path.join(base_dir, "Conversations - Cleaned")

    if not os.path.isdir(answers_dir):
        print(f"Error: The answers directory was not found at '{answers_dir}'")
        return
    if not os.path.isdir(conversations_dir):
        print(f"Error: The conversations directory was not found at '{conversations_dir}'")
        return

    students_data = {}

    # First, populate the students_data dictionary with answers and task types
    for filename in os.listdir(answers_dir):
        if filename.endswith('.docx') and not filename.startswith('._'):
            student_id = extract_student_id(filename)
            task_type = extract_task_type(filename)
            if student_id and task_type:
                file_path = os.path.join(answers_dir, filename)
                final_submission = process_answer_file(file_path)
                students_data[student_id] = {
                    "student_id": student_id,
                    "dialogue_history": [],
                    "final_submission": final_submission,
                    "task_type": task_type
                }

    # Now, process the conversation files and add the dialogue history
    for filename in os.listdir(conversations_dir):
        if filename.endswith('.docx') and not filename.startswith('._'):
            student_id = extract_student_id(filename)
            if student_id in students_data:
                file_path = os.path.join(conversations_dir, filename)
                conversation_history = parse_conversation_file(file_path)
                students_data[student_id]["dialogue_history"] = conversation_history

    # Convert dictionary to list for the final JSON output
    output_data = sorted(students_data.values(), key=lambda x: int(x['student_id']))

    output_file = os.path.join(base_dir, "gpt_student_conversations_and_results.json")
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=4, ensure_ascii=False)

    print(f"Processing complete. Output written to {output_file}")

if __name__ == "__main__":
    main() 